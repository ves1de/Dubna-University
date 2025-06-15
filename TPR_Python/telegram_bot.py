from telegram import Update
from telegram.ext import ApplicationBuilder, CommandHandler, ContextTypes, MessageHandler, filters
import spacy
from PyPDF2 import PdfReader
import docx

def extract_text_from_pdf(file_path):
    pdf_reader = PdfReader(file_path)
    text = ' '.join([page.extract_text() for page in pdf_reader.pages])
    return text


def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    text = ' '.join([paragraph.text for paragraph in doc.paragraphs])
    return text


# Токен бота, полученный от BotFather
TOKEN = "TOKEN"

# Загрузка модели spaCy
nlp = spacy.load("ru_sentiment_model")
    

# Функция обработки команды /start
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    await context.bot.send_message(chat_id=update.effective_chat.id,
                                   text=("Привет! Я могу определить сгенерированный контент.\n"
                                         "Через твой отправленный текст\n"
                                         "или через твой отправленный документ (.docx, .pdf)\n"
                                         "работа проделана, Трусовом Иваном"
                                         ))


# Функция обработки команды /help
async def help_command(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    await context.bot.send_message(chat_id=update.effective_chat.id,
                                   text="Я могу определить сгенерированнанный контент. Попробуйте написать мне что-нибудь!")


# Функция обработки любых сообщений
async def echo(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    text = update.message.text
    result = ""
    doc = nlp(text)
    if (doc.cats['POSITIVE'] > 0.5):
        result = f"Ваш текст является сгенерированным, уверенность: {100*(doc.cats['POSITIVE']):.0f}%"
    else:
        result = f"Ваш текст уникальнее нейросетевого, уверенность: {100*(doc.cats['NEGATIVE']):.0f}%"

    # Отправка результатов
    await context.bot.send_message(chat_id=update.effective_chat.id, text=f"{result}")


# Функция обработки файлов

async def handle_file(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    file_id = update.message.document.file_id
    file_info = await context.bot.get_file(file_id)
    text = ""
    try:
        # Скачиваем файл с правильным расширением
        if update.message.document.file_name.endswith(".pdf"):
            await file_info.download_to_drive(custom_path="temp_file.pdf")
            text = extract_text_from_pdf("temp_file.pdf")
        elif update.message.document.file_name.endswith(".docx"):
            await file_info.download_to_drive(custom_path="temp_file.docx")
            text = extract_text_from_docx("temp_file.docx")
        else:
            await context.bot.send_message(chat_id=update.effective_chat.id, text="Неподдерживаемый формат файла.")
            return

        # Обработка текста
        result = ""
        doc = nlp(text)
        if doc.cats['POSITIVE'] > 0.5:
            result = f"Ваш текст является сгенерированным, уверенность: {100*(doc.cats['POSITIVE']):.0f}%"
        else:
            result = f"Ваш текст уникальнее нейросетевого, уверенность: {100*(doc.cats['NEGATIVE']):.0f}%"

        await context.bot.send_message(chat_id=update.effective_chat.id, text=f"{result}")
    except Exception as e:
        await context.bot.send_message(chat_id=update.effective_chat.id, text=f"Ошибка при обработке файла: {str(e)}")


def main() -> None:
    # Создание приложения бота
    application = ApplicationBuilder().token(TOKEN).build()

    # Регистрация обработчиков команд
    application.add_handler(CommandHandler("start", start))
    application.add_handler(CommandHandler("help", help_command))

    # Регистрация обработчика сообщений
    application.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, echo))

    # Регистрация обработчика файлов
    application.add_handler(MessageHandler(filters.Document().ALL, handle_file))

    # Запуск бота
    application.run_polling()


if __name__ == "__main__":
    main()
